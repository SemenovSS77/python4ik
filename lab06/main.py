from appJar import gui
from recipe_calculator import get_recipe_ingredients, calculate_calories, calculate_cost, save_calculation, get_stats
from docx import Document
from openpyxl import Workbook
from datetime import datetime

def save_to_docx(recipe_name, calories, cost):
    doc = Document()
    doc.add_heading(f"Рецепт: {recipe_name}", 0)
    doc.add_paragraph(f"Энергетическая ценность: {calories:.2f} ккал")
    doc.add_paragraph(f"Стоимость: {cost:.2f} $")
    doc.add_paragraph(f"Дата расчета: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.save(f"{recipe_name}_report.docx")

def save_to_xlsx(recipe_name, calories, cost):
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет"
    
    # Заголовки
    ws['A1'] = "Рецепт"
    ws['B1'] = recipe_name
    ws['A2'] = "Энергетическая ценность"
    ws['B2'] = f"{calories:.2f} ккал"
    ws['A3'] = "Стоимость"
    ws['B3'] = f"{cost:.2f} $"
    ws['A4'] = "Дата расчета"
    ws['B4'] = datetime.now().strftime('%Y-%m-%d %H:%M')
    
    # Автоподбор ширины столбцов
    for column in ['A', 'B']:
        ws.column_dimensions[column].width = 20
    
    wb.save(f"{recipe_name}_report.xlsx")

def format_stats(stats):
    if not stats:
        return "Статистика отсутствует"
    
    result = ["Статистика расчетов:", "-------------------"]
    for name, count, avg_cal, avg_cost in stats:
        result.append(
            f"{name}: {count} расчётов\n"
            f"• Средние калории: {avg_cal:.1f} ккал\n"
            f"• Средняя стоимость: {avg_cost:.2f}$\n"
        )
    return "\n".join(result)

def on_calculate(button):
    recipe_name = app.getOptionBox("Рецепт")
    
    if not recipe_name:
        app.errorBox("Ошибка", "Выберите рецепт!")
        return
    
    if button in ["Рассчитать", "Сохранить в DOCX", "Сохранить в XLSX"]:
        # Расчет параметров
        ingredients = get_recipe_ingredients(recipe_name)
        calories = calculate_calories(ingredients)
        cost = calculate_cost(ingredients)
        
        # Обновление интерфейса
        app.setLabel("calories", f"Энергетическая ценность: {calories:.2f} ккал")
        app.setLabel("cost", f"Стоимость: {cost:.2f} $")
        
        # Сохранение в БД
        save_calculation(recipe_name, calories, cost)
        
        # Действия для конкретных кнопок
        if button == "Сохранить в DOCX":
            save_to_docx(recipe_name, calories, cost)
            app.infoBox("Успех", f"Отчет для {recipe_name} сохранен в DOCX!")
        elif button == "Сохранить в XLSX":
            save_to_xlsx(recipe_name, calories, cost)
            app.infoBox("Успех", f"Отчет для {recipe_name} сохранен в XLSX!")
    
    elif button == "Показать статистику":
        stats = get_stats()
        app.infoBox("Статистика расчетов", format_stats(stats))

# Создание GUI
app = gui("Калькулятор рецептов", "500x400")
app.setFont(size=12, family="Arial")

# Элементы интерфейса
app.addLabel("title", "Калькулятор параметров рецептов", 0, 0, 2)  # Заголовок
app.setLabelBg("title", "#e0e0e0")

app.addLabel("recipe_label", "Выберите рецепт:", 1, 0)
app.addOptionBox("Рецепт", ["Вок", "Бургер", "Пицца"], 1, 1)

# Кнопки действий
buttons = [
    "Рассчитать",
    "Сохранить в DOCX",
    "Сохранить в XLSX",
    "Показать статистику"
]
app.addButtons(buttons, on_calculate, 2, 0, 2)

# Области вывода результатов
app.addLabel("calories", "", 3, 0, 2)
app.addLabel("cost", "", 4, 0, 2)

# Стилизация
app.setButtonFont(size=11)
for button in buttons:
    app.setButtonBg(button, "#f0f0f0")

app.go()