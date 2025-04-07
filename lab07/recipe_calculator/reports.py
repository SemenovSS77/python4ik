from abc import ABC, abstractmethod

class Report(ABC):
    # Абстрактный базовый класс для отчетов
    @abstractmethod
    def save(self, recipe_name, calories, cost):
        pass

    def __str__(self):
        return f"Report: {self.__class__.__name__}"

    def __repr__(self):
        return f"<{self.__class__.__name__}>"


class DocxReport(Report): # Наследование от Report
    # Класс для создания отчетов в формате DOCX
    def save(self, recipe_name, calories, cost):
        from docx import Document
        doc = Document()
        doc.add_heading(f"Рецепт: {recipe_name}", 0)
        doc.add_paragraph(f"Энергетическая ценность: {calories:.2f} ккал")
        doc.add_paragraph(f"Стоимость: {cost:.2f} $")
        doc.save(f"{recipe_name}_report.docx")
        return True


class XlsxReport(Report): # Наследование от Report
    # Класс для создания отчетов в формате XLSX
    def save(self, recipe_name, calories, cost):
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Отчет"
        ws['A1'] = "Рецепт"
        ws['B1'] = recipe_name
        ws['A2'] = "Энергетическая ценность"
        ws['B2'] = calories
        ws['A3'] = "Стоимость"
        ws['B3'] = cost
        wb.save(f"{recipe_name}_report.xlsx")
        return True